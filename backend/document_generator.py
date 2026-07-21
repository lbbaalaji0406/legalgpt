"""
SAULGPT — DOCUMENT GENERATOR (.docx)
======================================
4-family architecture with LLM-generated structured specs.

Families: letter, pleading, affidavit, agreement
Sections: 10+1 rendering primitives
Guardrails: stamp margin, Triad, verification split, pagination signature,
            dispatch mode, annexure scanner
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import re, io
from typing import Optional

# ─────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────

DEFAULT_FONT  = "Times New Roman"
BODY_SIZE     = Pt(11)
HEADING_SIZE  = Pt(14)
TITLE_SIZE    = Pt(16)
MARGINS       = Cm(2.54)
STAMP_MARGIN  = Cm(14.0)  # ~5.5 inches for stamp paper

# ─────────────────────────────────────────────────────────────
# BASE HELPERS
# ─────────────────────────────────────────────────────────────

def _set_margins(doc: Document, stamp_paper: bool = False):
    for section in doc.sections:
        if stamp_paper:
            section.top_margin = STAMP_MARGIN
        else:
            section.top_margin = MARGINS
        section.bottom_margin = MARGINS
        section.left_margin   = MARGINS
        section.right_margin  = MARGINS


def _add_styled_paragraph(doc: Document, text: str, bold=False, size=BODY_SIZE,
                          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6),
                          font_name=DEFAULT_FONT, italic=False, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return p


def _add_body_paragraph(doc: Document, text: str, indent_first=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = BODY_SIZE
    return p


def _add_field_row(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(label + ":  ")
    run_label.font.name = DEFAULT_FONT
    run_label.font.size = BODY_SIZE
    run_label.bold = True
    run_value = p.add_run(value or "_________________")
    run_value.font.name = DEFAULT_FONT
    run_value.font.size = BODY_SIZE
    run_value.underline = not bool(value)
    return p


def _add_signature_block(doc: Document):
    doc.add_paragraph()
    _add_styled_paragraph(doc, "Yours faithfully,", space_after=Pt(24))
    _add_styled_paragraph(doc, "_________________", space_after=Pt(2))
    _add_styled_paragraph(doc, "(Signature)", space_after=Pt(2), size=Pt(10), italic=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Place:  _________________")
    r.font.name = DEFAULT_FONT; r.font.size = BODY_SIZE
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Date:   ")
    r2.font.name = DEFAULT_FONT; r2.font.size = BODY_SIZE
    r2.bold = True
    r2d = p2.add_run(date.today().strftime("%d %B %Y"))
    r2d.font.name = DEFAULT_FONT; r2d.font.size = BODY_SIZE


def _add_partner_signatures(doc: Document, party1_label="First Party", party2_label="Second Party"):
    """Two-party signature block for agreements."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r1 = p.add_run("_________________          _________________")
    r1.font.name = DEFAULT_FONT
    _add_styled_paragraph(doc, f"{party1_label:30s}      {party2_label:30s}",
                          align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9), italic=True)
    _add_styled_paragraph(doc, "(Sign here)                          (Sign here)",
                          align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9), italic=True)


def _add_witness_block(doc: Document):
    doc.add_paragraph()
    _add_styled_paragraph(doc, "In the presence of:", bold=True, space_after=Pt(8))
    for i in (1, 2):
        _add_styled_paragraph(doc, f"Witness {i}:", space_after=Pt(2))
        _add_styled_paragraph(doc, "Name:  ______________________", space_after=Pt(1), size=Pt(10))
        _add_styled_paragraph(doc, "Address:  ___________________", space_after=Pt(1), size=Pt(10))
        _add_styled_paragraph(doc, "Signature: __________________", space_after=Pt(8), size=Pt(10))


def _add_disclaimer(doc: Document):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("\u2014" * 40)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(128, 128, 128)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "This document is a draft prepared by SaulGPT for informational purposes only. "
        "It does not constitute legal advice. Please consult a qualified advocate "
        "before using this document for any legal proceeding."
    )
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(128, 128, 128)
    r2.font.italic = True


# ─────────────────────────────────────────────────────────────
# FIELD INTERPOLATION
# ─────────────────────────────────────────────────────────────

FIELD_PATTERN = re.compile(r"\{(\w+)\}")

def _interpolate(text: str, filled_fields: dict) -> str:
    """Replace {field_key} placeholders with actual values."""
    def replacer(m):
        key = m.group(1)
        return str(filled_fields.get(key, m.group(0)))
    return FIELD_PATTERN.sub(replacer, text)


# ─────────────────────────────────────────────────────────────
# ANNEXURE SCANNER
# ─────────────────────────────────────────────────────────────

ANNEXURE_PATTERN = re.compile(r"Annexure\s+([A-Z0-9])|Exhibit\s+([A-Z0-9])", re.IGNORECASE)

def _scan_annexures(sections: list, filled_fields: dict) -> list:
    """Extract annexure/exhibit references from interpolated section content."""
    refs = set()
    for sec in sections:
        text = str(sec)
        text = _interpolate(text, filled_fields)
        for match in ANNEXURE_PATTERN.finditer(text):
            refs.add(match.group(0))
    return sorted(refs)


# ─────────────────────────────────────────────────────────────
# 10+1 SECTION TYPE RENDERERS
# Each receives (doc, section_dict, filled_fields)
# ─────────────────────────────────────────────────────────────

def _render_heading(doc, sec, filled_fields):
    text = _interpolate(sec.get("text", ""), filled_fields)
    align = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
             "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(sec.get("alignment", "center"))
    _add_styled_paragraph(doc, text, bold=True, size=HEADING_SIZE, align=align)


def _render_subheading(doc, sec, filled_fields):
    text = _interpolate(sec.get("text", ""), filled_fields)
    align = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
             "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(sec.get("alignment", "center"))
    _add_styled_paragraph(doc, text, italic=True, size=Pt(10), align=align)


def _render_body_p(doc, sec, filled_fields):
    text = _interpolate(sec.get("text", ""), filled_fields)
    _add_body_paragraph(doc, text, indent_first=sec.get("indent", True))


def _render_numbered_list(doc, sec, filled_fields):
    items = sec.get("items", [])
    start = sec.get("start", 1)
    for i, item in enumerate(items, start):
        text = _interpolate(item, filled_fields)
        # Strip leading numbers from LLM content (e.g., "1. The Party..." → "The Party...")
        text = re.sub(r"^\s*\d+[\.\)]\s*", "", text)
        _add_body_paragraph(doc, f"{i}.  {text}")


def _render_field_row(doc, sec, filled_fields):
    label = _interpolate(sec.get("label", ""), filled_fields)
    value_field = sec.get("value_field")
    value = _interpolate(f"{{{value_field}}}", filled_fields) if value_field else sec.get("value", "")
    _add_field_row(doc, label, value)


def _render_schedule_box(doc, sec, filled_fields):
    """Property schedule with boundary directions in tabular format."""
    title = _interpolate(sec.get("title", "Schedule of Property"), filled_fields)
    _add_styled_paragraph(doc, title, bold=True, space_after=Pt(8),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    boundaries = sec.get("boundaries", {})
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    for i, (direction, desc) in enumerate(boundaries.items()):
        row = table.rows[i]
        row.cells[0].text = direction.upper()
        row.cells[1].text = _interpolate(str(desc), filled_fields)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = DEFAULT_FONT
                    run.font.size = Pt(10)
    remaining = sec.get("description", "")
    if remaining:
        doc.add_paragraph()
        _add_body_paragraph(doc, _interpolate(remaining, filled_fields))


def _render_signature(doc, sec, filled_fields):
    party1 = _interpolate(sec.get("party1_label", ""), filled_fields)
    party2 = _interpolate(sec.get("party2_label", ""), filled_fields)
    if party1 or party2:
        _add_partner_signatures(doc, party1, party2)
    else:
        _add_signature_block(doc)


def _render_prayer(doc, sec, filled_fields):
    items = sec.get("items", [])
    _add_styled_paragraph(doc, "P R A Y E R", bold=True, space_after=Pt(8),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_body_paragraph(doc, "In light of the above, it is most respectfully prayed that this Hon'ble Court may be pleased to:", indent_first=False)
    for i, item in enumerate(items, 1):
        text = _interpolate(item, filled_fields)
        _add_body_paragraph(doc, f"({chr(96 + i)})  {text}", indent_first=False)
    _add_body_paragraph(doc, "And pass any other order(s) that this Hon'ble Court may deem fit in the interests of justice.", indent_first=False)
    _add_styled_paragraph(doc, "Place:  _________________", space_after=Pt(2))
    _add_styled_paragraph(doc, f"Date:   {date.today().strftime('%d %B %Y')}", space_after=Pt(12))
    _add_styled_paragraph(doc, "Petitioner / Counsel for Petitioner", space_after=Pt(2))


def _render_attestation(doc, sec, filled_fields):
    """Notary/Oath Commissioner block for affidavits."""
    doc.add_paragraph()
    _add_styled_paragraph(doc, "VERIFICATION", bold=True, space_after=Pt(8),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    paras = sec.get("paragraphs", [])
    for p_text in paras:
        text = _interpolate(p_text, filled_fields)
        _add_body_paragraph(doc, text, indent_first=False)
    doc.add_paragraph()
    _add_styled_paragraph(doc, "Solemnly affirmed at _________________", space_after=Pt(4))
    _add_styled_paragraph(doc, f"on this {date.today().strftime('%d %B %Y')}", space_after=Pt(12))
    _add_styled_paragraph(doc, "IDENTIFIED BY ME,", bold=True, space_after=Pt(2))
    _add_styled_paragraph(doc, "[Advocate's Signature & Enrollment No.]", size=Pt(10), italic=True, space_after=Pt(12))
    _add_styled_paragraph(doc, "BEFORE ME,", bold=True, space_after=Pt(2))
    _add_styled_paragraph(doc, "[Notary / Oath Commissioner]", size=Pt(10), italic=True, space_after=Pt(4))
    _add_styled_paragraph(doc, "(Seal)", size=Pt(10), italic=True)


def _render_enclosure(doc, sec, filled_fields):
    items = sec.get("items", [])
    if not items:
        return
    doc.add_paragraph()
    _add_styled_paragraph(doc, "Enclosures:", bold=True, space_after=Pt(4))
    for item in items:
        _add_styled_paragraph(doc, f"  {_interpolate(item, filled_fields)}", space_after=Pt(2), size=Pt(10))


def _render_cc(doc, sec, filled_fields):
    items = sec.get("items", [])
    if not items:
        return
    doc.add_paragraph()
    _add_styled_paragraph(doc, "Cc:", bold=True, space_after=Pt(4))
    for item in items:
        _add_styled_paragraph(doc, f"  {_interpolate(item, filled_fields)}", space_after=Pt(2), size=Pt(10))


def _render_witness(doc, sec, filled_fields):
    _add_witness_block(doc)


# ─────────────────────────────────────────────────────────────
# SECTION DISPATCHER
# ─────────────────────────────────────────────────────────────

SECTION_RENDERERS = {
    "heading":       _render_heading,
    "subheading":    _render_subheading,
    "body_p":        _render_body_p,
    "numbered_list": _render_numbered_list,
    "field_row":     _render_field_row,
    "schedule_box":  _render_schedule_box,
    "signature":     _render_signature,
    "prayer":        _render_prayer,
    "attestation":   _render_attestation,
    "enclosure":     _render_enclosure,
    "cc":            _render_cc,
    "witness":       _render_witness,
}


# ─────────────────────────────────────────────────────────────
# GUARDRAIL VALIDATORS
# ─────────────────────────────────────────────────────────────

def validate_document_spec(spec: dict, filled_fields: dict = None) -> dict:
    """
    Validate a document spec against all guardrails.
    Returns {"valid": bool, "errors": [str], "warnings": [str]}.
    """
    import json
    errors = []
    warnings = []
    family = spec.get("family", "generic_instrument")
    sections = spec.get("sections", [])
    section_types = {s.get("type") for s in sections}

    # ── Universal: at least one body section ──
    if not sections:
        errors.append("Spec has no sections.")

    # ── Universal: check for prompt leakage in body_p ──
    PROMPT_LEAK_PATTERNS = [
        "i need to draft", "i need a", "i want to draft", "write a", "prepare a",
        "the user", "user's description", "user query", "the client said",
        "i need to create", "i am drafting", "this notice is to inform you that the user",
    ]
    for sec in sections:
        if sec.get("type") == "body_p":
            text = sec.get("text", "").lower()
            if any(p in text for p in PROMPT_LEAK_PATTERNS):
                errors.append(f"Prompt leakage detected in body_p: '{sec.get('text', '')[:80]}...'")

    # ── LETTER family ──
    if family == "letter":
        if not spec.get("dispatch_mode"):
            errors.append("Letters require a dispatch_mode (RPAD, Speed Post, Email, etc.)")

        # Check for statutory demand clause in body_p
        demand_indicators = ["call upon you", "hereby", "pay the sum", "within.*days",
                            "failing which", "shall be constrained", "legal proceedings",
                            "without further notice", "initiate", "recovery"]
        has_demand = any(
            any(re.search(p, _interpolate(s.get("text", ""), filled_fields or {}), re.IGNORECASE)
                for p in demand_indicators)
            for s in sections if s.get("type") in ("body_p", "numbered_list")
        )
        if not has_demand:
            errors.append("Letter is missing a statutory demand clause with deadline and legal consequences")

        # Check for address content
        address_present = any(
            "address" in _interpolate(s.get("text", ""), filled_fields or {}).lower()
            or any(k in (filled_fields or {}) for k in ["sender_full_address", "recipient_full_address"])
            for s in sections if s.get("type") in ("body_p", "field_row")
        )
        if not address_present:
            warnings.append("Letter may be missing full party addresses")

    # ── PLEADING family: Triad of Survival ──
    if family == "pleading":
        triad_indicators = ["limitation", "cause of action arose", "period of limitation"]
        has_limitation = any(
            any(ind in _interpolate(s.get("text", "") + str(s.get("items", "")), filled_fields or {}).lower()
                for ind in triad_indicators)
            for s in sections if s.get("type") in ("body_p", "numbered_list")
        )
        if not has_limitation:
            errors.append("Pleading is missing a limitation clause (cause of action / period of limitation)")

        jurisdiction_indicators = ["jurisdiction", "territorial", "this hon'ble court"]
        has_jurisdiction = any(
            any(ind in _interpolate(s.get("text", "") + str(s.get("items", "")), filled_fields or {}).lower()
                for ind in jurisdiction_indicators)
            for s in sections if s.get("type") in ("body_p", "numbered_list")
        )
        if not has_jurisdiction:
            errors.append("Pleading is missing a jurisdiction clause")

        valuation_indicators = ["valuation", "court fee", "pecuniary", "suits valuation"]
        has_valuation = any(
            any(ind in _interpolate(s.get("text", "") + str(s.get("items", "")), filled_fields or {}).lower()
                for ind in valuation_indicators)
            for s in sections if s.get("type") in ("body_p", "numbered_list")
        )
        if not has_valuation:
            warnings.append("Pleading may be missing a valuation / court fee clause")

        if not has_limitation or not has_jurisdiction:
            errors.append("Pleading rejected: Triad of Survival incomplete")

    # ── AFFIDAVIT family: verification split ──
    if family == "affidavit":
        has_attestation = "attestation" in section_types
        has_signature = "signature" in section_types
        has_body = any(s.get("type") == "body_p" for s in sections)
        if not has_attestation:
            errors.append("Affidavit missing attestation/verification block")
        if not has_body:
            errors.append("Affidavit has no body content")

    # ── AGREEMENT family: pagination handled by renderer ──
    if family == "agreement":
        has_witness = "witness" in section_types or "signature" in section_types
        if not has_witness:
            errors.append("Agreement missing signature/witness block")

    # ── Annexure scanner ──
    annexures = _scan_annexures(
        [json.dumps(s) if isinstance(s, dict) else str(s) for s in sections],
        filled_fields or {}
    ) if filled_fields else []
    has_enclosure = "enclosure" in section_types
    if annexures and not has_enclosure:
        warnings.append(f"Body references annexures/exhibits ({', '.join(annexures)}) but no Enclosures section found")

    return {"valid": len(errors) == 0, "errors": errors, "warnings": warnings}


# ─────────────────────────────────────────────────────────────
# 4 FAMILY BUILDERS + GENERIC
# ─────────────────────────────────────────────────────────────

def _build_letter(doc: Document, spec: dict, filled_fields: dict):
    """One-way communication with dispatch mode, To/From, Subject, body."""
    stamp = spec.get("requires_stamp_paper", False)
    _set_margins(doc, stamp)

    # ── Hardcoded RPAD header ──
    _add_styled_paragraph(doc, "REGISTERED POST WITH ACKNOWLEDGEMENT DUE (RPAD)",
                          bold=True, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(12))

    # ── Hardcoded TO: block ──
    recipient_name = filled_fields.get("recipient_name", filled_fields.get("recipient_full_name", ""))
    recipient_addr = filled_fields.get("recipient_address", filled_fields.get("recipient_full_address", ""))
    if recipient_name:
        _add_styled_paragraph(doc, "TO,", bold=True, space_after=Pt(4))
        to_text = recipient_name
        if recipient_addr:
            to_text += f"\n{recipient_addr}"
        _add_body_paragraph(doc, to_text, indent_first=False)
        doc.add_paragraph()

    sections = spec.get("sections", [])
    for sec in sections:
        renderer = SECTION_RENDERERS.get(sec.get("type"))
        if renderer:
            renderer(doc, sec, filled_fields)

    # ── Hardcoded Advocate Signature Footer ──
    _add_styled_paragraph(doc, "Yours faithfully,", space_after=Pt(24))
    sender = filled_fields.get("sender_name", filled_fields.get("sender_full_name", ""))
    _add_styled_paragraph(doc, sender if sender else "_________________",
                          bold=True, space_after=Pt(2))
    _add_styled_paragraph(doc, "Advocate", size=Pt(10), italic=True, space_after=Pt(2))
    _add_styled_paragraph(doc, "Enrollment No.: _________________", size=Pt(10), space_after=Pt(4))
    _add_disclaimer(doc)


def _build_pleading(doc: Document, spec: dict, filled_fields: dict):
    """Court filing — Cause title, Parties, Narrative body_p, Prayer, Verification."""
    _set_margins(doc)
    jurisdiction = spec.get("sub_jurisdiction", "DISTRICT_CIVIL")

    # ── Hardcoded Cause Title ──
    cause_title = {
        "ORIGINAL_SIDE_HC": "IN THE HIGH COURT OF JUDICATURE AT _________________\nORIGINAL CIVIL JURISDICTION",
        "DISTRICT_CIVIL": "IN THE COURT OF THE DISTRICT JUDGE / CIVIL JUDGE (SR. DIV.) AT _________________",
        "CRIMINAL_MAGISTRATE": "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE / METROPOLITAN MAGISTRATE AT _________________",
        "TRIBUNAL_NCLT": "IN THE NATIONAL COMPANY LAW TRIBUNAL, _________________ BENCH",
        "CONSUMER_FORUM": "IN THE CONSUMER DISPUTES REDRESSAL COMMISSION / FORUM AT _________________",
    }.get(jurisdiction, "IN THE COURT OF _________________ AT _________________")
    _add_styled_paragraph(doc, cause_title, bold=True, size=Pt(12),
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))
    _add_styled_paragraph(doc, "SUIT / PETITION / COMPLAINT NO. _________ OF _________",
                          size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    # ── Hardcoded Parties block ──
    parties = spec.get("parties", {})
    plaintiff_field = parties.get("plaintiff", "plaintiff_name")
    defendant_field = parties.get("defendant", "defendant_name")
    pl_name = filled_fields.get(plaintiff_field, "_________________")
    de_name = filled_fields.get(defendant_field, "_________________")
    pl_addr = filled_fields.get("plaintiff_address", filled_fields.get("sender_full_address", ""))
    de_addr = filled_fields.get("defendant_address", filled_fields.get("recipient_full_address", ""))
    _add_field_row(doc, "Petitioner / Plaintiff", pl_name)
    if pl_addr:
        _add_field_row(doc, "Address", pl_addr)
    _add_field_row(doc, "Respondent / Defendant", de_name)
    if de_addr:
        _add_field_row(doc, "Address", de_addr)
    doc.add_paragraph()

    # ── LLM-generated content: body_p, prayer, field_row (skip attestation, rendered below) ──
    sections = spec.get("sections", [])

    # Count actual numbered_list items — Python len(), not LLM hallucination
    numbered_sections = [s for s in sections if s.get("type") == "numbered_list"]
    total_paras = len(numbered_sections[0].get("items", [])) if numbered_sections else 0
    knowledge_end = total_paras - 1 if total_paras > 1 else 1
    filled_fields["knowledge_range"] = f"1 to {knowledge_end}"
    filled_fields["advice_index"] = str(total_paras)
    filled_fields["date"] = date.today().strftime('%d %B %Y')

    for sec in sections:
        if sec.get("type") == "attestation":
            continue
        renderer = SECTION_RENDERERS.get(sec.get("type"))
        if renderer:
            renderer(doc, sec, filled_fields)

    # ── Dynamic Verification Block (same page as signature) ──
    doc.add_paragraph()
    ver_heading = _add_styled_paragraph(doc, "VERIFICATION", bold=True, space_after=Pt(8),
                                        align=WD_ALIGN_PARAGRAPH.CENTER)
    ver_heading.paragraph_format.keep_with_next = True
    att_para = ""
    for sec in sections:
        if sec.get("type") == "attestation":
            paras = sec.get("paragraphs", [])
            if paras:
                att_para = _interpolate(paras[0], filled_fields)
            break
    if not att_para:
        att_para = (
            "I, the above-named Petitioner / Plaintiff, do hereby verify that the contents of "
            f"paragraphs {filled_fields.get('knowledge_range', '1 to 3')} of the accompanying petition "
            f"are true to my personal knowledge and that paragraph "
            f"{filled_fields.get('advice_index', '4 to 5')} is based on legal advice which I believe "
            "to be true. Nothing material has been concealed. Verified at _________________ "
            f"on this {date.today().strftime('%d %B %Y')}."
        )
    ver_para = _add_body_paragraph(doc, att_para, indent_first=False)
    ver_para.paragraph_format.keep_with_next = True
    sig = _add_styled_paragraph(doc, "Petitioner / Plaintiff", space_after=Pt(4))
    sig.paragraph_format.keep_with_next = True
    _add_disclaimer(doc)


def _build_affidavit(doc: Document, spec: dict, filled_fields: dict):
    """Sworn statement — Stamp margin, Oath heading, Knowledge body_p, Verification, Identification."""
    stamp = spec.get("requires_stamp_paper", True)
    _set_margins(doc, stamp)
    title = spec.get("display_name", "AFFIDAVIT")
    _add_styled_paragraph(doc, title.upper(), bold=True, size=HEADING_SIZE,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── Hardcoded Oath Heading ──
    deponent = filled_fields.get("deponent_name", filled_fields.get("sender_full_name", "_________________"))
    father = filled_fields.get("father_name", "_________________")
    age = filled_fields.get("deponent_age", "_________________")
    addr = filled_fields.get("deponent_address", filled_fields.get("sender_full_address", "_________________"))
    oath_text = (
        f"I, {deponent}, son / daughter of {father}, aged about {age}, "
        f"residing at {addr}, do hereby solemnly affirm and state as follows:"
    )
    _add_body_paragraph(doc, oath_text, indent_first=True)
    doc.add_paragraph()

    # ── LLM-generated content: body_p (numbered paragraphs), attestation ──
    sections = spec.get("sections", [])

    # Count actual numbered_list items BEFORE rendering — Python len() never lies
    numbered_sections = [s for s in sections if s.get("type") == "numbered_list"]
    total_paras = len(numbered_sections[0].get("items", [])) if numbered_sections else 0
    knowledge_end = total_paras - 1 if total_paras > 1 else 1
    filled_fields["knowledge_range"] = f"1 to {knowledge_end}"
    filled_fields["advice_index"] = str(total_paras)
    filled_fields["city"] = filled_fields.get("city", "Chennai")
    filled_fields["date"] = date.today().strftime('%d %B %Y')

    for sec in sections:
        if sec.get("type") == "attestation":
            # Render verification in the hardcoded block below instead
            continue
        renderer = SECTION_RENDERERS.get(sec.get("type"))
        if renderer:
            renderer(doc, sec, filled_fields)

    # ── Hardcoded Strict Verification Clause ──
    # Read paragraph counts from attestation section in spec
    att_para = ""
    for sec in sections:
        if sec.get("type") == "attestation":
            paras = sec.get("paragraphs", [])
            if paras:
                att_para = _interpolate(paras[0], filled_fields)
            break
    if not att_para:
        att_para = (
            f"Verified at _________________ on this {date.today().strftime('%d %B %Y')}, "
            f"that the contents of paragraphs 1 to 3 are true to my personal knowledge, "
            f"and paragraphs 4 to 5 are based on legal advice which I believe to be true. "
            f"Nothing material has been concealed."
        )
    doc.add_paragraph()
    ver_heading = _add_styled_paragraph(doc, "VERIFICATION", bold=True, space_after=Pt(8),
                                        align=WD_ALIGN_PARAGRAPH.CENTER)
    ver_heading.paragraph_format.keep_with_next = True
    ver_para = _add_body_paragraph(doc, att_para, indent_first=False)
    ver_para.paragraph_format.keep_with_next = True

    # ── Hardcoded Identification Block (same page as verification) ──
    id_advocate = _add_styled_paragraph(doc, "IDENTIFIED BY ME,", bold=True, space_after=Pt(2))
    id_advocate.paragraph_format.keep_with_next = True
    _add_styled_paragraph(doc, "Advocate:  _________________", space_after=Pt(2), size=Pt(10))
    _add_styled_paragraph(doc, "Enrollment No.:  _________________", space_after=Pt(12), size=Pt(10))
    id_notary = _add_styled_paragraph(doc, "BEFORE ME,", bold=True, space_after=Pt(2))
    id_notary.paragraph_format.keep_with_next = True
    _add_styled_paragraph(doc, "[Notary / Oath Commissioner]", space_after=Pt(4), size=Pt(10), italic=True)
    _add_styled_paragraph(doc, "(Seal)", space_after=Pt(4), size=Pt(10), italic=True)
    _add_disclaimer(doc)


def _build_agreement(doc: Document, spec: dict, filled_fields: dict):
    """Bilateral instrument — Stamp margin, Party definitions, WHEREAS recitals, Clauses, Schedule, Dispute resolution, Witness."""
    stamp = spec.get("requires_stamp_paper", True)
    _set_margins(doc, stamp)
    title = spec.get("display_name", "AGREEMENT")
    _add_styled_paragraph(doc, title.upper(), bold=True, size=HEADING_SIZE,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    _add_styled_paragraph(doc, "Made on this _____ day of ________________, 20____",
                          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ── Hardcoded Party Definitions (direct from filled_fields, bypass spec mapping) ──
    party1_name = filled_fields.get("party1_name", "First Party")
    party2_name = filled_fields.get("party2_name", "Second Party")
    party1_addr = filled_fields.get("party1_address", filled_fields.get("sender_full_address", ""))
    party2_addr = filled_fields.get("party2_address", filled_fields.get("recipient_full_address", ""))
    _add_styled_paragraph(doc, "BETWEEN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    party1_def = f"{party1_name}, having its registered office at {party1_addr or '_________________'}"
    party2_def = f"{party2_name}, having its registered office at {party2_addr or '_________________'}"

    pan1 = filled_fields.get("party1_pan", "")
    pan2 = filled_fields.get("party2_pan", "")
    if pan1:
        party1_def += f" (PAN: {pan1})"
    if pan2:
        party2_def += f" (PAN: {pan2})"

    _add_body_paragraph(doc,
        f"{party1_def} (hereinafter called the 'First Party' which expression shall "
        f"include their heirs, successors, and assigns) of the ONE PART.")
    doc.add_paragraph()
    _add_styled_paragraph(doc, "AND", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    _add_body_paragraph(doc,
        f"{party2_def} (hereinafter called the 'Second Party' which expression shall "
        f"include their heirs, successors, and assigns) of the OTHER PART.")
    doc.add_paragraph()

    # ── LLM-generated content: body_p (WHEREAS recitals), numbered_list (operative clauses), schedule_box ──
    sections = spec.get("sections", [])
    for sec in sections:
        renderer = SECTION_RENDERERS.get(sec.get("type"))
        if renderer:
            renderer(doc, sec, filled_fields)

    # ── No hardcoded dispute resolution block — LLM's own clause stands alone ──

    # ── Hardcoded Witness Block ──
    _add_witness_block(doc)

    # ── Pagination Signature Footer ──
    from docx.oxml.ns import qn
    for i, section in enumerate(doc.sections):
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("_________________  (Party Initials)")
        r.font.name = DEFAULT_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(128, 128, 128)

    _add_disclaimer(doc)


def _build_generic(doc: Document, spec: dict, filled_fields: dict):
    """Fallback for unrecognized types. Strict structural safety net."""
    stamp = spec.get("requires_stamp_paper", False)
    _set_margins(doc, stamp)
    title = spec.get("display_name", "LEGAL INSTRUMENT")
    _add_styled_paragraph(doc, title.upper(), bold=True, size=HEADING_SIZE,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _add_styled_paragraph(doc, f"Date: {date.today().strftime('%d %B %Y')}",
                          space_after=Pt(12))
    sections = spec.get("sections", [])
    for sec in sections:
        renderer = SECTION_RENDERERS.get(sec.get("type"))
        if renderer:
            renderer(doc, sec, filled_fields)
    if "signature" not in {s.get("type") for s in sections}:
        _add_signature_block(doc)
    _add_disclaimer(doc)


# ─────────────────────────────────────────────────────────────
# FAMILY DISPATCHER
# ─────────────────────────────────────────────────────────────

FAMILY_BUILDERS = {
    "letter":             _build_letter,
    "pleading":           _build_pleading,
    "affidavit":          _build_affidavit,
    "agreement":          _build_agreement,
    "generic_instrument": _build_generic,
}


# ─────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────

def generate_docx(
    display_name: str,
    filled_fields: dict,
    doc_spec: Optional[dict] = None,
    doc_type: Optional[str] = None,
    is_dynamic: bool = False,
    problem_description: str = "",
) -> bytes:
    """
    Generate a professionally formatted .docx from a document spec.

    Args:
        display_name: Human-readable document name
        filled_fields: Dict of field_key → user-provided value
        doc_spec: Structured spec dict with family + sections (NEW entry point)
        doc_type: Legacy parameter — kept for backward compat
        is_dynamic: Legacy parameter — kept for backward compat
        problem_description: Original user problem text

    Returns:
        .docx file contents as bytes
    """
    doc = Document()

    if doc_spec:
        # ── NEW PATH: spec-based generation ──
        builder = FAMILY_BUILDERS.get(doc_spec.get("family", "generic_instrument"), _build_generic)
        builder(doc, doc_spec, filled_fields)
    else:
        # ── LEGACY PATH: for backward compat during transition ──
        # Will be removed once triage spec flow is fully wired
        from docx import Document as _D
        from docx.shared import Pt, Cm
        _set_margins(doc)
        _add_styled_paragraph(doc, (display_name or "LEGAL DOCUMENT").upper(),
                              bold=True, size=TITLE_SIZE,
                              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        for key, val in filled_fields.items():
            if val:
                _add_field_row(doc, key.replace("_", " ").title(), val)
        if problem_description:
            doc.add_paragraph()
            _add_body_paragraph(doc, problem_description)
        _add_signature_block(doc)
        _add_disclaimer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
